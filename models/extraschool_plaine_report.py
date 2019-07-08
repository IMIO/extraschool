# -*- coding: utf-8 -*-
##############################################################################
#
#    Extraschool
#    Copyright (C) 2008-2014
#    Jean-Michel Abé - Town of La Bruyère (<http://www.labruyere.be>)
#    Michael Michot & Michael Colicchial - Imio (<http://www.imio.be>).
#
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU Affero General Public License as
#    published by the Free Software Foundation, either version 3 of the
#    License, or (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU Affero General Public License for more details.
#
#    You should have received a copy of the GNU Affero General Public License
#    along with this program.  If not, see <http://www.gnu.org/licenses/>.
#
##############################################################################
from docutils.utils.math.math2html import Formula

from openerp import models, api, fields,_
import base64
import os
import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from datetime import date, datetime, timedelta as td
from openerp.exceptions import except_orm, Warning, RedirectWarning
import io
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
import docx
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.text.tabstops import TabStop as ts
from docx.text.parfmt import ParagraphFormat
from docx.shared import Pt, Inches


class extraschool_plaine_report(models.Model):
    _name = 'extraschool.plaine_report'

    name = fields.Char('Nom', required=True)
    validity_from = fields.Date('Date from')
    validity_to = fields.Date('Date to',)
    placeid = fields.Many2one('extraschool.place')
    activity_ids = fields.Many2many(
        'extraschool.activity',
        'extraschool_activity_plaine_rel',
        'plaine_id',
        'activity_id',
        string='Activity',

    )
    guardian_ids = fields.Many2many(
        'extraschool.guardian',
        'extraschool_guardian_plaine_rel',
        'plaine_id',
        'guardian_id',
        string='Guardian',

    )

    @api.multi
    def generate_report(self, plaine_report):
        Po_obj = self.env['extraschool.organising_power'].search([])[0]
        # Select the path of the new object 'document'
        path = '/tmp/plaine_report' + str(datetime.now()) + '.docx'
        document = Document()

        # Start of generating document.
        # Start Page 1

        font = document.styles['Normal'].font
        font.name = 'Arial Narrow'
        font.size = Pt(9)
        font.color.rgb = RGBColor(0x0, 0x0, 0x0)

        title = document.add_heading(u'Centres de Vacances', 2)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.bold = True
        title.center = True

        title2 = document.add_paragraph()
        title2.add_run(u'Formulaire de demande de subventionnement').bold = True
        title2.underline = True
        title2.size = Pt(13)
        title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        PO_titre = document.add_heading(u'POUVOIR ORGANISATEUR', 9)
        PO_titre.underline = True
        PO_titre.bold = True

        po_tel = Po_obj.po_tel or ''
        po_email = Po_obj.po_email or ''
        po_zipcode = Po_obj.po_zipcode or ''
        po_city = Po_obj.po_city or ''
        po_street = Po_obj.po_street or ''
        po_town = Po_obj.town or ''

        PO_paragraph = document.add_paragraph()
        PO_paragraph.add_run(u'Dénomination : ' + po_town + '\n'
                         u'Adresse : ' + po_street + u'\t C.P.: ' + po_zipcode + u'\t Ville/Commune : ' + po_city + '\n'
                         u' Tel : '+ po_tel + u'\t Fax : \t Courriel : ' + po_email ).bold = True

        PO_case = document.add_heading(u'Pouvoir Public \t Organisation de jeunesse reconnue \t Autre.',9)
        PO_case.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        CV_titre = document.add_heading(u'CENTRE DE VACANCES', 9)
        CV_titre.underline = True
        CV_titre.bold = True

        CV_case = document.add_heading(u'Plaine de vacances \t Séjour de vacances \t Camp de vacances \n'
                                       u'Pour les séjours et les camps : \t Infrastructures résidentielles \t Sous tente', 9)
        CV_case.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cv_name = plaine_report['placeid'].name
        cv_street = plaine_report['placeid'].street or ''
        cv_zipcode = plaine_report['placeid'].zipcode or ''
        cv_city = plaine_report['placeid'].city or ''
        cv_email = plaine_report['placeid'].email or ''
        cv_tel = plaine_report['placeid'].tel or ''

        CV_paragraph = document.add_paragraph()
        CV_paragraph.add_run(u'Nom du centre (le cas échéant) : ' +  cv_name + '\n'
                                              u'Adresse : ' + cv_street + u'\t C.P.: ' + cv_zipcode + u'\t Ville/Commune : '+ cv_city +  u'\n'
                                              u'Tel : ' + cv_tel + u'\t Fax : \t Courriel : ' + cv_email).bold = True



        CO_titre = document.add_heading(u'CORRESPONDANT', 9)
        CO_titre.underline = True
        CO_titre.bold = True

        CO_paragraph = document.add_paragraph()
        CO_paragraph.add_run(u'Nom : \t Prénom : \n'
                                              u'Fonction : \n'
                                              u'Adresse : \t C.P. : \t Ville/Commune : \n'
                                              u'Tel : \t Fax : \t Courriel :').bold = True
        CO_paragraph.bold = True

        CF_titre = document.add_heading(u'COMPTE FINANCIER', 9)
        CF_titre.underline = True
        CF_titre.bold = True

        CF_paragraph = document.add_paragraph()
        CF_paragraph.add_run(u'N° de compte : \n'
                                              u'Titulaire : ' + po_town + '\n'
                                              u'Adresse : ' + po_street + '\n'
                                              u'Code Postal : ' + po_zipcode + u'\t Ville/Commune : ' + cv_city).bold = True
        CF_paragraph.bold = True

        Remarque_titre = document.add_heading(u'REMARQUE IMPORTANTE', 9)
        Remarque_titre.underline = True

        texte_page1 = document.add_paragraph(u'Pour être recevable, la présente demande doit impérativement comprendre :')
        texte_page1_2 = document.add_paragraph(u'la liste des enfants accueillis ;',style = 'List Bullet')
        texte_page1_3 = document.add_paragraph(
            u'la liste du personnel d\'encadrement (animateurs, coodinateur(s), responsable qualifié, animateurs ou'
            u'coordinateur(s) en 2ème stage de formation), accompagnée des demandes d\'assimilation s\'il s\'agit'
            u'de la première prestation de l\'animateur ou du coordinateur concerné ;', style = 'List Bullet')
        texte_page1_4 = document.add_paragraph(u'le tableau de présences journalières (enfants et animateurs) ;',style = 'List Bullet')
        texte_page1_5 = document.add_paragraph(
            u'le cas échéant, le justificatif de versement des indémnités aux animateurs et coordinateurs, sous la '
            u'forme d\'une déclaration sur l\'honneur.',style = 'List Bullet')

        paragraphe_page_1 = document.add_paragraph()
        paragraphe_page_1.add_run(u'Sans ces annexes complètes, il ne peut être procédé à l\'examen du droit à la subvention ni au calcul de celle-ci !').bold = True

        paragraphe2_page_1 = document.add_paragraph()
        paragraphe2_page_1.add_run(u'Ce formulaire est à transmettre à l\'O.N.E au plus tard le 30 septembre pour les vacances \n'
                                                    u'd\'été ou 30 jours après la fin des activités pour les vacances de Noël ou de Pâques.'
                                                    u'Ce formulaire est indispensable au calcul des subventions méritées par chaque centre de vacances !').bold = True
        paragraphe2_page_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER





        # records = (
        #     (3, 'Dénomination', '', ''),
        #     (7, 'Adresse', 'C.P.:', 'Ville/Commune :'),
        #     (4, 'Tel', 'Fax:', 'Courriel')
        # )
        #
        # table = document.add_table(rows=1, cols=4)
        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'POUVOIR ORGANISATEUR   '
        # hdr_cells[1].text = 'j'
        # hdr_cells[2].text = 'j'
        # hdr_cells[2].text = ' '
        # for qty, id, desc, desc2 in records:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = qty
        #     row_cells[1].text = id
        #     row_cells[2].text = desc
        #     row_cells[2].text = desc2


        document.add_page_break()

        # End of generation. First save document, then export to attachment and remove file from /tmp
        document.save(path)

        outfile = open(path, "r").read()
        attachment_obj = self.env['ir.attachment']
        attachment_obj.create({'res_model': 'extraschool.plaine_report', 'res_id': plaine_report.id,
                               'datas_fname': 'rapport_plaine_' + '.docx',
                               'name': 'rapport_plaine_' + '.docx',
                               'datas': base64.b64encode(outfile), })

        os.remove(path)


    @api.model
    def create(self, vals):

        id = super(extraschool_plaine_report, self).create(vals)
        self.generate_report(id)

        return id


        '''report = '/tmp/plaine_report' + str(datetime.now()) + '.xls'
        workbook = xlsxwriter.Workbook(report)

        title = workbook.add_format({'bold': True, 'bg_color': '#cccccc', 'font_size': 20, 'border': True})
        bold = workbook.add_format({'bold': True, 'border': True})
        border = workbook.add_format({'border': True})

        # - de 6 ans
        worksheet = workbook.add_worksheet('- 6 ans')
        worksheet.merge_range('A1:E1', "ENFANTS DE - DE 6 ANS", title)
        worksheet.write(1,0, "Nom", bold)
        worksheet.write(1,1, u"Prénom", bold)
        worksheet.write(1,2, "Age", bold)
        worksheet.write(1,3, "Nombre de jours", bold)
        worksheet.write(1,4, "Prix", bold)

        row = 2
        total_prestation=total_price = 0

        for child in under_6:
            worksheet.write(row, 0, under_6[child]['lastname'], border)
            worksheet.write(row, 1, under_6[child]['firstname'], border)
            worksheet.write(row, 2, str(under_6[child]['age']) + " ans", border)
            worksheet.write(row, 3, under_6[child]['prestation'], border)
            worksheet.write(row, 4, str(under_6[child]['price']) + "€".decode('utf-8'), border)
            row += 1
            total_prestation += under_6[child]['prestation']
            total_price += under_6[child]['price']
        worksheet.merge_range('A' + str(row +1) + ':C' + str(row +1), "Total", bold)
        worksheet.write(row, 3, total_prestation, border)
        worksheet.write(row, 4, str(total_price) + "€".decode('utf-8'), border)

        # + de 6 ans
        worksheet = workbook.add_worksheet('+ 6 ans')
        worksheet.merge_range('A1:E1', "ENFANTS DE + DE 6 ANS", title)
        worksheet.write(1,0, "Nom", bold)
        worksheet.write(1,1, u"Prénom", bold)
        worksheet.write(1,2, "Age", bold)
        worksheet.write(1,3, "Nombre de jours", bold)
        worksheet.write(1,4, "Prix", bold)

        row = 2
        total_prestation = total_price = 0

        for child in over_6:
            worksheet.write(row, 0, over_6[child]['lastname'], border)
            worksheet.write(row, 1, over_6[child]['firstname'], border)
            worksheet.write(row, 2, str(over_6[child]['age']) + " ans", border)
            worksheet.write(row, 3, over_6[child]['prestation'], border)
            worksheet.write(row, 4, str(over_6[child]['price']) + "€".decode('utf-8'), border)
            row += 1
            total_prestation += over_6[child]['prestation']
            total_price += over_6[child]['price']
        worksheet.merge_range('A' + str(row + 1) + ':C' + str(row + 1), "Total", bold)
        worksheet.write(row, 3, total_prestation, border)
        worksheet.write(row, 4, str(total_price) + "€".decode('utf-8'), border)

        workbook.close()

        outfile = open(report, "r").read()
        attachment_obj = self.env['ir.attachment']
        attachment_obj.create({'res_model': 'extraschool.plaine_report', 'res_id': id.id,
                               'datas_fname': 'rapport_plaine_' + '.xls',
                               'name': 'rapport_one_' + '.xls',
                               'datas': base64.b64encode(outfile), })
        os.remove(report)

    @api.model
    def create(self,vals):

        under_6 = {}
        over_6 = {}

        prestation_ids = self.env['extraschool.prestationtimes'].search([
            ('prestation_date', '>=', vals.get('start_date')),
            ('prestation_date', '<=', vals.get('end_date')),
            ('activity_occurrence_id.activityid', 'in', tuple(vals.get('activity')[0][2])),
            ('es', '=', 'E'),
        ]).sorted(key=lambda r: (r.childid.lastname))

        for prestation in prestation_ids :
            age = prestation.childid.get_age()
            if age >= 6 :
                if prestation.childid.id in over_6:
                    over_6[prestation.childid.id]['prestation'] += 1
                    over_6[prestation.childid.id]['price'] += prestation.invoiced_prestation_id.total_price
                else:
                    over_6[prestation.childid.id] = {'lastname': prestation.childid.lastname.upper(),
                                                     'firstname': prestation.childid.firstname, 'age': age,
                                                     'prestation': 1,
                                                     'price': prestation.invoiced_prestation_id.total_price}
            else:
                if prestation.childid.id in under_6:
                    under_6[prestation.childid.id]['prestation'] += 1
                    under_6[prestation.childid.id]['price'] += prestation.invoiced_prestation_id.total_price
                else:
                    under_6[prestation.childid.id] = {'lastname': prestation.childid.lastname.upper(),
                                                      'firstname': prestation.childid.firstname, 'age': age,
                                                      'prestation': 1,
                                                      'price': prestation.invoiced_prestation_id.total_price}

        id = super(extraschool_plaine_report, self).create(vals)
        self.generate_report(under_6,over_6,id)

        return id'''
