# -*- coding: utf-8 -*-
##############################################################################
#
#    Extraschool
#    Copyright (C) 2008-2019
#    Jean-Michel Abé - Town of La Bruyère (<http://www.labruyere.be>)
#    Michael Michot & Michael Colicchia & Jenny Pans - Imio (<http://www.imio.be>).
#
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU Affero General Public License as
#    published by the Free Software Foundation, either version 3 of the
#    License, or (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU Affero General Public License for more details.
#
#    You should have received a copy of the GNU Affero General Public License
#    along with this program.  If not, see <http://www.gnu.org/licenses/>.
#
##############################################################################
import base64
import os
from datetime import datetime
from docx import Document

from openerp import models, api, fields


class extraschool_plain_report(models.Model):
    _name = 'extraschool.plain_report'
    _description = 'Plain report'

    name = fields.Char('Nom', required=True)
    start_date = fields.Date('Start date', required=True)
    end_date = fields.Date('End date', required=True)
    dates_ok = fields.Boolean(default=True)
    dates_warning = fields.Boolean(default=False)
    activity_ids = fields.Many2many(
        'extraschool.activity',
        'extraschool_activity_activity_plain_rel',
        'plain_report_id',
        'plain_activity_id',
        string='Activity',
        required=True)
    title = fields.Selection(
        (('public_power', 'Pouvoir public'),
         ('organisation', 'Organisation de jeunesse reconnue'),
         ('other', 'Autre')),
        default='public_power', string='Title')
    camps_radio = fields.Selection(
        (('holiday_plain', 'Plaine de vacances'),
         ('stays', u'Séjour de vacances'),
         ('holiday_camp', 'Camp de vacances'),
         ('residential_infra', 'Infrastructures Résidentielles'),
         ('tent', 'Sous tente')),
        default='holiday_plain', string='Stays and camps')

    @api.onchange('start_date', 'end_date')
    @api.multi
    def _check_validity_date(self):
        for record in self:
            if record.start_date and record.end_date:
                if record.end_date < record.start_date:
                    record.dates_ok = False
                    record.dates_warning = True
                else:
                    record.dates_ok = True
                    record.dates_warning = False

    @api.multi
    def _generate_subvention_request(self, document):
        """
        Generate a subvention request.
        """
        sections = [
            ['POUVOIR ORGANISATEUR\n', 'Dénomination : ', 'Adresse : ', 'C.P : ', 'Ville/Commune : ', 'Tel :', 'Fax :', 'Courriel :'],
            ['CENTRE DE VACANCES\n', 'Nom du centre (le cas échéant) : ', 'Adresse : ', 'C.P : ', 'Ville/Commune : ', 'Tel :', 'Fax :', 'Courriel :'],
            ['CORRESPONDANT\n', 'NOM : ', 'PRENOM : ', 'Fonction : ', 'Adresse : ', 'C.P : ', 'Ville/Commune : ', 'Tel :', 'Fax :', 'Courriel :'],
            ['COMPTE FINANCIER\n', 'N° de compte : ', 'Titulaire : ', 'Adresse : ', 'C.P : ', 'Ville/Commune : ']
        ]
        document.add_heading('Centre de Vacances', 0)
        document.add_heading('Formulaire de demande de subventionnement', 1)
        for section in sections:
            document.add_heading(section[0], 2)
            paragraph = document.add_paragraph()
            for i in range(1, len(section)):
                paragraph.add_run(section[i].decode('utf-8')).bold = True

        document.add_heading(u'REMARQUE IMPORTANTE')
        document.add_paragraph(u'Pour être recevable, la présente demande doit impérativement comprendre :')
        document.add_paragraph(u'la liste des enfants accueillis ; ').style = 'List Bullet'
        document.add_paragraph(u'la liste du personnel d’encadrement (animateurs, coordinateur(s), responsable qualifié, animateurs ou coordinateur(s) en 2ème stage de formation), accompagnée des demandes d’assimilation s’il s’agit de la première prestation de l’animateur ou du coordinateur concerné ; ').style = 'List Bullet'
        document.add_paragraph(u'le tableau de présences journalières (enfants et animateurs) ; ').style = 'List Bullet'
        document.add_paragraph(u'e cas échéant, le justificatif de versement des indemnités aux animateurs et coordinateurs, sous la forme d’une déclaration sur l’honneur. ').style = 'List Bullet'
        document.add_paragraph().add_run(u'Sans ces annexes complètes, il ne peut être procédé à l’examen du droit à la subvention ni au calcul de celle-ci ! ').bold = True
        paragraph = document.add_paragraph()
        paragraph.add_run(u'Ce formulaire est à transmettre à l’O.N.E. au plus tard le 30 septembre pour les vacances d’été ou 30 jours après la fin des activités pour les vacances de Noël ou de Pâques. Ce formulaire est indispensable au calcul des subventions méritées par chaque centre de vacances!').bold = True

    @api.multi
    def _generate_summary(self, document):
        pass

    @api.multi
    def _generate_childs_list(self, document):
        pass

    @api.multi
    def _generate_supervisors(self, document):
        pass

    @api.multi
    def _generate_prestation_times(self, document):
        pass

    @api.multi
    def generate_report(self, id):
        report = '/tmp/plain_report' + str(datetime.now()) + '.docx'
        document = Document()
        self._generate_subvention_request(document)
        self._generate_summary(document)
        self._generate_childs_list(document)
        self._generate_supervisors(document)
        self._generate_prestation_times(document)
        document.save(report)
        outfile = open(report, "r").read()
        attachment_obj = self.env['ir.attachment']
        attachment_obj.create({'res_model': 'extraschool.plain_report', 'res_id': id.id,
                               'datas_fname': 'rapport_subvention_' + '.docx',
                               'name': 'plain_report_' + '.docx',
                               'datas': base64.b64encode(outfile), })
        os.remove(report)

    @api.model
    def create(self, vals):
        id = super(extraschool_plain_report, self).create(vals)
        self.generate_report(id)
        return id
